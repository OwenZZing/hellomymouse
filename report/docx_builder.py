"""Build the Research Starter Kit Word document from JSON data."""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from report.templates import (
    FONT_NAME, COLOR_PRIMARY, COLOR_SECONDARY, COLOR_MUTED,
    COLOR_WARN, COLOR_GREEN_MSG, COLOR_GREEN_TEXT, COLOR_TREND,
    COLOR_STARS_HIGH, COLOR_STARS_LOW, COLOR_ATTRIBUTION, COLOR_LEGEND, COLOR_TIME_NOTE,
    BG_ALT, FEAS_COLORS, FEAS_FC, COST_COLORS, STAR_LEGEND, ENCOURAGE_MSG,
    COLOR_EASY_FC, COLOR_MEDIUM_FC,
)


# ── Low-level helpers ─────────────────────────────────────────

def sf(run, bold=False, size=10, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1, color=None):
    if color is None:
        color = COLOR_PRIMARY
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    sz = 16 if level == 1 else 13 if level == 2 else 11
    sf(run, bold=True, size=sz, color=color)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    return p


def body(doc, text, color=None):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    sf(run, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    sf(run, color=color)
    return p


def tbl_header(table, bg=None):
    if bg is None:
        bg = (0, 70, 127)
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*bg))
        tcPr.append(shd)
        for para in cell.paragraphs:
            for r in para.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.name = FONT_NAME
                r.font.size = Pt(10)


def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def tbl_style(table):
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(9)


# ── Document builder ──────────────────────────────────────────

def build_report(data: dict, output_path: str, review: dict | None = None) -> str:
    """Build the Word document from Stage 2 JSON data. Returns output_path."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    lab_name = data.get('lab_name', 'Research Lab')
    field = data.get('field', '')
    lab_overview = data.get('lab_overview', '')
    assigned_project = data.get('assigned_project', '')
    projects = data.get('projects', [])
    hypotheses = data.get('hypotheses', [])
    paper_summaries = data.get('paper_summaries', [])
    capabilities = data.get('lab_capabilities', {})
    costs = data.get('costs', [])
    checklist = data.get('checklist', [])
    bg_knowledge = data.get('background_knowledge', {})
    roadmap = data.get('roadmap', [])
    intro = data.get('intro_for_undergrad', {})

    # ── Cover ────────────────────────────────────────────────
    _build_cover(doc, lab_name, field, review=review)

    # ── About page ───────────────────────────────────────────
    _build_about_page(doc)

    # ── TOC ─────────────────────────────────────────────────
    _build_toc(doc, capabilities)

    # ── Section 0-A: 연구란 무엇인가 (학부생 안내) ─────────────
    if intro:
        _build_intro_section(doc, intro)

    # ── Section 0-B: Projects ────────────────────────────────
    _build_section0(doc, projects, lab_overview, assigned_project)

    # ── Section 1: Lab Capabilities ──────────────────────────
    _build_section1(doc, capabilities)

    # ── Section 2: Paper summaries ───────────────────────────
    _build_section2(doc, paper_summaries)

    # ── Section 3: Hypotheses ────────────────────────────────
    _build_section3(doc, hypotheses, assigned_project, projects)

    # ── Section 4: Costs ─────────────────────────────────────
    _build_section4(doc, costs)

    # ── Section 5: PI Checklist ──────────────────────────────
    _build_section5(doc, checklist)

    # ── Section 6: Background Knowledge ─────────────────────
    _build_section6(doc, bg_knowledge)

    # ── Section 7: Roadmap ───────────────────────────────────
    _build_section7(doc, roadmap)

    doc.save(output_path)
    return output_path


# ── Section builders ─────────────────────────────────────────

def _build_cover(doc, lab_name, field, review: dict | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Research Starter Kit')
    sf(run, bold=True, size=28, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(lab_name)
    sf(run, bold=True, size=18, color=COLOR_SECONDARY)

    if field:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(field)
        sf(run, size=12, color=COLOR_MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n[주의] 이 리포트는 AI가 논문을 분석하여 자동 생성한 초안입니다.\n반드시 교수님 및 선배 연구원의 검토를 받으세요.\n\n')
    sf(run, size=9, color=COLOR_WARN)

    # ── Review testimonial on cover ───────────────────────────
    if review and any(review.values()):
        name    = review.get('name', '').strip()
        fld     = review.get('field', '').strip()
        stars   = review.get('stars', 0)
        comment = review.get('comment', '').strip()

        if comment or stars:
            # 구분선
            sep = doc.add_paragraph()
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = sep.add_run('─' * 28)
            sf(r, size=9, color=COLOR_ATTRIBUTION)

            # 별점 + 한줄평
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if stars:
                star_str = '★' * stars + '☆' * (5 - stars)
                c = COLOR_STARS_HIGH if stars >= 4 else COLOR_STARS_LOW
                r = p.add_run(f'{star_str}  ')
                sf(r, bold=True, size=11, color=c)
            if comment:
                r = p.add_run(f'"{comment}"')
                sf(r, size=10, color=COLOR_SECONDARY)
            p.paragraph_format.space_after = Pt(2)

            # 이름 + 분야
            if name or fld:
                byline = '— ' + ('  ·  '.join(filter(None, [name, fld])))
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p2.add_run(byline)
                sf(r, size=9, color=COLOR_MUTED)
                p2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Made by @hellomymouse  ·  kby930@gmail.com')
    sf(run, size=8, color=COLOR_ATTRIBUTION)

    doc.add_page_break()


def _build_about_page(doc):
    p = doc.add_paragraph()
    run = p.add_run('이 리포트에 대하여')
    sf(run, bold=True, size=18, color=COLOR_PRIMARY)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('만든 이유')
    sf(run, bold=True, size=12, color=COLOR_PRIMARY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    body(doc, '대학원에 처음 들어가면 교수님 면담 전 무엇을 공부해야 하는지, 어떤 연구를 제안해야 하는지 막막합니다. "연구실 논문을 읽어봐라"는 말은 들었지만 논문 수십 편을 며칠 안에 소화하고 아이디어까지 내기는 쉽지 않습니다. 이 프로그램은 그 첫 번째 벽을 낮추기 위해 만들었습니다.')

    p = doc.add_paragraph()
    run = p.add_run('목표')
    sf(run, bold=True, size=12, color=COLOR_PRIMARY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    bullet(doc, '연구실 논문을 자동 분석하여 "지금 이 연구실이 무엇을 하는가"를 빠르게 파악')
    bullet(doc, '기존 연구의 빈틈(Gap)에서 실현 가능한 가설 후보 도출')
    bullet(doc, '입학 첫 면담, 랩 미팅 발표, 연구 방향 결정에 바로 활용할 수 있는 초안 제공')

    p = doc.add_paragraph()
    run = p.add_run('누구를 위한 리포트인가')
    sf(run, bold=True, size=12, color=COLOR_PRIMARY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    bullet(doc, '연구실에 막 합류한 석사 / 박사 신입생')
    bullet(doc, '새 분야에 처음 들어가는 포닥 또는 연구 인턴')
    bullet(doc, '지도교수님 면담 전 연구 방향을 잡고 싶은 대학원생 누구나')

    p = doc.add_paragraph()
    run = p.add_run('사용 방법')
    sf(run, bold=True, size=12, color=COLOR_PRIMARY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    for step, text in [
        ('Step 1', '연구실의 최근 5년 논문 PDF를 최대한 모아 지정 폴더에 넣는다'),
        ('Step 2', 'AI가 논문을 훑어보며 프로젝트 목록을 자동 파악한다 (Stage 0)'),
        ('Step 3', '관심 있는 프로젝트를 선택하거나 전체 분석을 진행한다'),
        ('Step 4', '심층 분석이 완료되면 이 리포트가 자동으로 생성된다'),
        ('Step 5', '목차 순서대로 읽으며 교수님 면담 및 첫 랩미팅을 준비한다'),
    ]:
        _p = doc.add_paragraph()
        _r1 = _p.add_run(f'{step}   ')
        sf(_r1, bold=True, size=10, color=COLOR_PRIMARY)
        _r2 = _p.add_run(text)
        sf(_r2, size=10)
        _p.paragraph_format.left_indent = Pt(12)
        _p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run('처음 시작하는 여러분에게')
    sf(run, bold=True, size=12, color=COLOR_GREEN_MSG)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _p = doc.add_paragraph()
    _r = _p.add_run(ENCOURAGE_MSG)
    sf(_r, size=10, color=COLOR_GREEN_TEXT)
    _p.paragraph_format.left_indent = Pt(12)
    _p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('⚠  주의사항')
    sf(run, bold=True, size=12, color=COLOR_WARN)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for warn in [
        'AI가 논문 텍스트를 분석한 초안입니다. 내용이 부정확하거나 누락될 수 있습니다.',
        '가설의 실현 가능성 · 임팩트 별점은 AI의 추정치입니다. 선배 연구원 및 교수님의 검토가 반드시 필요합니다.',
        '비용 추정치는 참고용이며, 실제 lab 보유 장비 및 환경에 따라 크게 달라질 수 있습니다.',
        '이 리포트를 교수님께 제출할 경우 AI 생성 초안임을 반드시 명시하고, 스스로 검토한 후 활용하세요.',
    ]:
        _p = doc.add_paragraph()
        _r = _p.add_run(f'•  {warn}')
        sf(_r, size=9, color=COLOR_WARN)
        _p.paragraph_format.left_indent = Pt(12)
        _p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()


def _build_toc(doc, capabilities):
    p = doc.add_paragraph()
    run = p.add_run('목  차')
    sf(run, bold=True, size=18, color=COLOR_PRIMARY)
    p.paragraph_format.space_after = Pt(16)

    techniques = capabilities.get('techniques', [])
    equipment = capabilities.get('equipment_or_models', [])
    sub1 = '1-1. Core Techniques'
    sub2 = '1-2. Equipment & Models' if equipment else '1-2. Key Resources'

    for num, title, is_sub in [
        ('서문.', '연구란 무엇인가? — 처음 연구실에 온 당신을 위한 안내', False),
        ('',   '서문-1. 연구란 무엇인가?', True),
        ('',   '서문-2. 이 연구실은 무슨 연구를 하나요?', True),
        ('',   '서문-3. 왜 이 연구가 중요한가요?', True),
        ('',   '서문-4. 어떻게 연구하나요?', True),
        ('0.', '연구실 프로젝트 파악 결과 (Stage 0)', False),
        ('1.', '연구실 보유 기술 및 장비 (Lab Capabilities)', False),
        ('',   sub1, True),
        ('',   sub2, True),
        ('2.', '논문별 연구 요약', False),
        ('3.', '연구 가설 (Research Hypotheses)', False),
        ('',   '3-1. Feasibility Matrix', True),
        ('',   '3-2. Hypothesis 상세', True),
        ('4.', '소요 자원 및 비용 추정', False),
        ('5.', 'PI 확인 체크리스트', False),
        ('6.', '배경 지식 가이드 (Background Knowledge)', False),
        ('',   '6-1. 알아야 할 핵심 개념', True),
        ('',   '6-2. 추천 검색 키워드', True),
        ('',   '6-3. 추천 저널 및 학회', True),
        ('7.', '추천 첫 3개월 Roadmap', False),
    ]:
        _p = doc.add_paragraph()
        if is_sub:
            _r = _p.add_run(f'        {title}')
            sf(_r, size=10, color=(80, 80, 80))
            _p.paragraph_format.space_after = Pt(2)
        else:
            _r1 = _p.add_run(f'{num}  ')
            sf(_r1, bold=True, size=12, color=COLOR_PRIMARY)
            _r2 = _p.add_run(title)
            sf(_r2, bold=True, size=12, color=(30, 30, 30))
            _p.paragraph_format.space_before = Pt(6)
            _p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def _feas_bg(feas: str) -> str:
    """Fuzzy match feasibility → background hex color."""
    f = feas.lower()
    if 'easy' in f or 'fast' in f or 'low' in f: return 'C6EFCE'
    if 'hard' in f or 'explor' in f or 'high' in f: return 'FFC7CE'
    if 'medium' in f or 'moderate' in f or 'mid' in f: return 'FFEB9C'
    return 'FFEB9C'  # default to yellow (Medium) rather than white


def _feas_fc(feas: str) -> tuple:
    """Fuzzy match feasibility → font color RGB."""
    f = feas.lower()
    if 'easy' in f or 'fast' in f or 'low' in f: return (0, 128, 0)
    if 'hard' in f or 'explor' in f or 'high' in f: return (180, 0, 0)
    if 'medium' in f or 'moderate' in f or 'mid' in f: return (180, 130, 0)
    return (180, 130, 0)  # default yellow


def _build_intro_section(doc, intro: dict):
    """서문: 연구란 무엇인가 — 학부생을 위한 안내."""
    COLOR_GREEN = (60, 100, 60)

    heading(doc, '서문. 연구란 무엇인가? — 처음 연구실에 온 당신을 위한 안내', level=1, color=COLOR_GREEN)

    # 서문-1
    heading(doc, '서문-1. 연구(Research)란 무엇인가요?', level=2, color=COLOR_GREEN)
    what_is = intro.get('what_is_research', '')
    if what_is:
        body(doc, what_is)
    p = doc.add_paragraph()
    r = p.add_run('좋은 연구의 조건: (1) 아직 아무도 해보지 않은 질문, '
                  '(2) 그 질문에 답할 수 있는 실험/구현, '
                  '(3) 다른 사람이 재현할 수 있는 명확한 방법, '
                  '(4) 솔직한 결과 보고.')
    sf(r, size=10, color=(80, 80, 80))
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(8)

    # 서문-2
    heading(doc, '서문-2. 이 연구실은 무슨 연구를 하나요?', level=2, color=COLOR_GREEN)
    what_lab = intro.get('what_does_lab_do', '')
    if what_lab:
        body(doc, what_lab)

    # 서문-3
    heading(doc, '서문-3. 왜 이 연구가 중요한가요?', level=2, color=COLOR_GREEN)
    why = intro.get('why_important', '')
    if why:
        body(doc, why)

    # 서문-4
    heading(doc, '서문-4. 어떻게 연구하나요? (실험 방법 쉽게 보기)', level=2, color=COLOR_GREEN)
    methods = intro.get('how_they_study', [])
    if methods:
        body(doc, '이 연구실의 대표적인 연구 방법을 쉽게 설명하면:')
        for m in methods:
            mp = doc.add_paragraph()
            mp.paragraph_format.left_indent = Pt(14)
            mp.paragraph_format.space_after = Pt(4)
            mr1 = mp.add_run(f'▶ {m.get("method", "")}: ')
            sf(mr1, bold=True, size=10, color=COLOR_PRIMARY)
            mr2 = mp.add_run(m.get('plain_explanation', ''))
            sf(mr2, size=10)

    doc.add_paragraph()


def _build_section0(doc, projects, lab_overview, assigned_project):
    heading(doc, '0. Lab Projects 파악 결과 (Stage 0)', level=1)
    body(doc, lab_overview or '논문 분석으로 파악한 주요 연구 프로젝트입니다.')

    if assigned_project:
        p = doc.add_paragraph()
        r1 = p.add_run('📌 배정된 프로젝트: ')
        sf(r1, bold=True, size=10, color=COLOR_PRIMARY)
        r2 = p.add_run(assigned_project)
        sf(r2, bold=True, size=10, color=COLOR_TREND)
        p.paragraph_format.space_after = Pt(6)

    if not projects:
        body(doc, '(프로젝트 목록 없음)')
        doc.add_paragraph()
        return

    t = doc.add_table(rows=len(projects) + 1, cols=3)
    t.style = 'Table Grid'
    for j, h in enumerate(['#', 'Project', '관련 논문']):
        t.rows[0].cells[j].text = h
    tbl_header(t)
    for i, proj in enumerate(projects):
        row = t.rows[i + 1].cells
        row[0].text = str(proj.get('id', i + 1))
        name = proj.get('name', '')
        desc = proj.get('description', '')
        row[1].text = f'{name}\n{desc}' if desc else name
        papers = proj.get('related_papers', [])
        row[2].text = '\n'.join(papers) if papers else ''
        if i % 2 == 0:
            for cell in row:
                cell_bg(cell, BG_ALT)
    tbl_style(t)
    doc.add_paragraph()


def _build_section1(doc, capabilities):
    heading(doc, '1. Lab Capabilities', level=1)

    techniques = capabilities.get('techniques', [])
    equipment = capabilities.get('equipment_or_models', [])

    heading(doc, '1-1. Core Techniques', level=2)
    if techniques:
        t = doc.add_table(rows=len(techniques) + 1, cols=3)
        t.style = 'Table Grid'
        for j, h in enumerate(['Technique', '한 줄 설명', 'Lab 구축 현황']):
            t.rows[0].cells[j].text = h
        tbl_header(t)
        for i, tech in enumerate(techniques):
            row = t.rows[i + 1].cells
            row[0].text = tech.get('name', '')
            row[1].text = tech.get('description', '')
            row[2].text = tech.get('lab_status', '')
            if i % 2 == 0:
                for cell in row:
                    cell_bg(cell, BG_ALT)
        tbl_style(t)
    else:
        body(doc, '(기술 목록 없음)')
    doc.add_paragraph()

    heading(doc, '1-2. Equipment & Models', level=2)
    if equipment:
        t = doc.add_table(rows=len(equipment) + 1, cols=3)
        t.style = 'Table Grid'
        for j, h in enumerate(['Name', '한 줄 설명', '특징']):
            t.rows[0].cells[j].text = h
        tbl_header(t)
        for i, eq in enumerate(equipment):
            row = t.rows[i + 1].cells
            row[0].text = eq.get('name', '')
            row[1].text = eq.get('description', '')
            row[2].text = eq.get('notes', '')
            if i % 2 == 0:
                for cell in row:
                    cell_bg(cell, BG_ALT)
        tbl_style(t)
    else:
        body(doc, '(장비/모델 목록 없음)')
    doc.add_paragraph()


def _build_section2(doc, paper_summaries):
    heading(doc, '2. 논문별 연구 요약', level=1)
    if not paper_summaries:
        body(doc, '(논문 요약 없음)')
        return

    t = doc.add_table(rows=len(paper_summaries) + 1, cols=6)
    t.style = 'Table Grid'
    for j, h in enumerate(['논문', 'Tag', '유형', '요약', '핵심 결과', 'Limitation']):
        t.rows[0].cells[j].text = h
    tbl_header(t)
    type_colors = {
        'simulation': 'DDEBF7',
        'experimental': 'E2EFDA',
        'theoretical': 'FFF2CC',
        'review': 'F2F2F2',
        'mixed': 'FCE4D6',
    }
    for i, s in enumerate(paper_summaries):
        row = t.rows[i + 1].cells
        title = s.get('title', s.get('filename', ''))
        ptype = s.get('paper_type', '').lower().split()[0] if s.get('paper_type') else ''
        row[0].text = title[:60] + ('...' if len(title) > 60 else '')
        row[1].text = s.get('method_tag', '')
        row[2].text = ptype if ptype else '-'
        row[3].text = s.get('summary', '')[:120]
        row[4].text = s.get('key_finding', '')[:100]
        row[5].text = s.get('limitation', '')[:100]
        if ptype in type_colors:
            cell_bg(row[2], type_colors[ptype])
        if i % 2 == 0:
            for ci in [0, 1, 3, 4, 5]:
                cell_bg(row[ci], BG_ALT)
    tbl_style(t)
    doc.add_paragraph()


def _build_section3(doc, hypotheses, assigned_project, projects):
    heading(doc, '3. Research Hypotheses', level=1)

    if assigned_project:
        body(doc, f'배정된 프로젝트 ({assigned_project}) 관련 가설을 우선 배치했습니다. H1~H6은 기존 논문 Limitation에서 직접 파생된 안전한 가설, H7은 2023~2025 최신 연구 트렌드를 접목한 신박한 가설입니다.')
    else:
        body(doc, 'H1~H6은 기존 논문 Limitation에서 직접 파생된 안전한 가설, H7은 2023~2025 최신 연구 트렌드를 접목한 신박한 가설입니다. 교수님께는 H1~H3 중 하나 + H7 조합으로 제안하면 좋은 반응을 얻을 가능성이 높습니다.')

    heading(doc, '3-1. Feasibility Matrix', level=2)

    # Legend
    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(4)
    for stars, desc in STAR_LEGEND:
        _r = legend.add_run(f'  {stars} {desc}  ')
        sf(_r, size=8, color=COLOR_LEGEND)

    # Learning time note
    _note = doc.add_paragraph()
    _nr = _note.add_run('※ 예상 기간은 해당 기술을 이미 숙지한 경우 기준입니다. 신기술 학습 곡선을 포함하면 실제로는 2~3배 소요될 수 있습니다.')
    sf(_nr, size=8, color=COLOR_TIME_NOTE)
    _note.paragraph_format.space_before = Pt(2)
    _note.paragraph_format.space_after = Pt(6)

    if not hypotheses:
        body(doc, '(가설 없음)')
        return

    t = doc.add_table(rows=len(hypotheses) + 1, cols=5)
    t.style = 'Table Grid'
    for j, h in enumerate(['ID', 'Hypothesis', 'Feasibility', '논문 임팩트 / 기간', '핵심 아이디어 (요약)']):
        t.rows[0].cells[j].text = h
    tbl_header(t)

    for i, hypo in enumerate(hypotheses):
        row = t.rows[i + 1].cells
        hid = hypo.get('id', f'H{i+1}')
        feas = hypo.get('feasibility', 'Medium')
        impact_stars = hypo.get('impact_stars', '★★★☆☆')
        impact_desc = hypo.get('impact_desc', '')
        period = hypo.get('period', '')

        row[0].text = hid
        row[1].text = hypo.get('name', '')
        row[2].text = feas
        row[3].text = f'{impact_stars}\n{period}\n{impact_desc}'
        stmt = hypo.get('statement', '')
        row[4].text = stmt[:85] + ('...' if len(stmt) > 85 else '')

        cell_bg(row[2], _feas_bg(feas))
        if hid == 'H7':
            for cell in [row[0], row[1], row[3], row[4]]:
                cell_bg(cell, 'FFF2CC')

    tbl_style(t)
    doc.add_paragraph()

    # 3-2 Detail
    heading(doc, '3-2. Hypothesis 상세', level=2)
    feas_fc = FEAS_FC

    for hypo in hypotheses:
        hid = hypo.get('id', '')
        name = hypo.get('name', '')
        feas = hypo.get('feasibility', 'Medium')
        period = hypo.get('period', '')
        impact_stars = hypo.get('impact_stars', '★★★☆☆')

        p = doc.add_paragraph()
        color = COLOR_TREND if hid == 'H7' else COLOR_PRIMARY
        run1 = p.add_run(f'{hid}. {name}  ')
        sf(run1, bold=True, size=11, color=color)
        run2 = p.add_run(f'[{feas} / {period}]')
        sf(run2, bold=True, size=10, color=_feas_fc(feas))
        run_imp = p.add_run(f'  {impact_stars}')
        sf(run_imp, size=10, color=COLOR_STARS_HIGH if '★★★★' in impact_stars else COLOR_STARS_LOW)
        if hid == 'H7':
            run3 = p.add_run('  ★ 최신 트렌드 접목')
            sf(run3, bold=True, size=10, color=(200, 100, 0))
        p.paragraph_format.space_before = Pt(10)

        bullet(doc, f'Hypothesis: {hypo.get("statement", "")}')

        novelty = hypo.get('novelty', '')
        if novelty:
            bullet(doc, f'Novelty: {novelty}')

        bullet(doc, f'Research Gap & 근거: {hypo.get("rationale", "")}')

        # Evaluation metrics
        metrics = hypo.get('evaluation_metrics', [])
        baseline = hypo.get('baseline', '')
        if metrics or baseline:
            _p = doc.add_paragraph(style='List Bullet')
            _r = _p.add_run('Evaluation: ')
            sf(_r, bold=True, size=9, color=COLOR_PRIMARY)
            if baseline:
                _r2 = _p.add_run(f'Baseline — {baseline}')
                sf(_r2, size=9)
            if metrics:
                for m in metrics:
                    _mp = doc.add_paragraph(style='List Bullet')
                    _mp.paragraph_format.left_indent = Pt(24)
                    metric_text = m.get('metric', '')
                    target = m.get('target', '')
                    method = m.get('measurement_method', '')
                    _mr = _mp.add_run(f'{metric_text}  →  {target}  ({method})')
                    sf(_mr, size=9, color=(60, 60, 60))

        # Control loop design
        ctrl = hypo.get('control_loop_design', '')
        if ctrl and ctrl.strip().upper() != 'N/A':
            _p = doc.add_paragraph(style='List Bullet')
            _r1 = _p.add_run('제어 루프 설계: ')
            sf(_r1, bold=True, size=9, color=(30, 90, 160))
            _r2 = _p.add_run(ctrl)
            sf(_r2, size=9, color=(30, 60, 120))

        # Fallback plan
        fallback = hypo.get('fallback_plan', '')
        if fallback:
            _p = doc.add_paragraph(style='List Bullet')
            _r1 = _p.add_run('Plan B: ')
            sf(_r1, bold=True, size=9, color=(150, 80, 0))
            _r2 = _p.add_run(fallback)
            sf(_r2, size=9, color=(100, 60, 0))

        bullet(doc, f'Required resources: {hypo.get("resources", "")}')

    doc.add_paragraph()


def _build_section4(doc, costs):
    heading(doc, '4. 소요 자원 및 비용 추정 (KRW 기준)', level=1)
    body(doc, '* 비용은 참고 수치입니다. 실제 lab 보유 재료 및 환경에 따라 달라집니다.')
    if not costs:
        body(doc, '(비용 목록 없음)')
        doc.add_paragraph()
        return

    t = doc.add_table(rows=len(costs) + 1, cols=4)
    t.style = 'Table Grid'
    for j, h in enumerate(['항목', '비용 범주', '추정 비용 (KRW)', '비고']):
        t.rows[0].cells[j].text = h
    tbl_header(t)
    from report.templates import COST_COLORS
    for i, cost in enumerate(costs):
        row = t.rows[i + 1].cells
        row[0].text = cost.get('item', '')
        cat = cost.get('category', 'Medium')
        row[1].text = cat
        row[2].text = cost.get('estimated_krw', '')
        row[3].text = cost.get('note', '')
        cell_bg(row[1], COST_COLORS.get(cat, 'FFFFFF'))
        if i % 2 == 0:
            cell_bg(row[0], BG_ALT)
            cell_bg(row[2], BG_ALT)
            cell_bg(row[3], BG_ALT)
    tbl_style(t)
    doc.add_paragraph()


def _build_section5(doc, checklist):
    heading(doc, '5. PI 확인 체크리스트 (MANDATORY)', level=1)
    body(doc, '분석 시작 전 반드시 교수님 및 선배와 확인하세요.')
    for item in checklist:
        bullet(doc, f'☐  {item}')
    doc.add_paragraph()


def _build_section6(doc, bg_knowledge):
    heading(doc, '6. Background Knowledge Guide', level=1)

    heading(doc, '6-1. 알아야 할 핵심 개념', level=2)
    concepts = bg_knowledge.get('core_concepts', [])
    if concepts:
        t = doc.add_table(rows=len(concepts) + 1, cols=3)
        t.style = 'Table Grid'
        for j, h in enumerate(['개념', '설명', '왜 필요한가']):
            t.rows[0].cells[j].text = h
        tbl_header(t)
        for i, c in enumerate(concepts):
            row = t.rows[i + 1].cells
            row[0].text = c.get('concept', '')
            row[1].text = c.get('description', '')
            row[2].text = c.get('why_needed', '')
            if i % 2 == 0:
                for cell in row:
                    cell_bg(cell, BG_ALT)
        tbl_style(t)
    else:
        body(doc, '(핵심 개념 없음)')
    doc.add_paragraph()

    heading(doc, '6-2. 추천 검색 키워드', level=2)
    keywords = bg_knowledge.get('search_keywords', [])
    if keywords:
        body(doc, '  ·  '.join(keywords))
    doc.add_paragraph()

    heading(doc, '6-3. 추천 저널 및 학회', level=2)
    journals = bg_knowledge.get('recommended_journals', [])
    if journals:
        t = doc.add_table(rows=len(journals) + 1, cols=3)
        t.style = 'Table Grid'
        for j, h in enumerate(['저널 / 학회', '분야', '추천 이유']):
            t.rows[0].cells[j].text = h
        tbl_header(t)
        for i, j in enumerate(journals):
            row = t.rows[i + 1].cells
            row[0].text = j.get('name', '')
            row[1].text = j.get('field', '')
            row[2].text = j.get('why', '')
            if i % 2 == 0:
                for cell in row:
                    cell_bg(cell, BG_ALT)
        tbl_style(t)
    else:
        body(doc, '(저널 목록 없음)')
    doc.add_paragraph()


def _build_section7(doc, roadmap):
    heading(doc, '7. 추천 첫 3개월 Roadmap', level=1)
    if not roadmap:
        body(doc, '(Roadmap 없음)')
        return
    for month_data in roadmap:
        period = month_data.get('period', '')
        tasks = month_data.get('tasks', [])
        p = doc.add_paragraph()
        run = p.add_run(period)
        sf(run, bold=True, size=11, color=COLOR_PRIMARY)
        p.paragraph_format.space_before = Pt(8)
        for task in tasks:
            bullet(doc, task)
    doc.add_paragraph()


def _build_review_section(doc, review: dict):
    """사용자 리뷰 섹션 (선택사항)."""
    doc.add_page_break()
    heading(doc, '리뷰', level=1, color=COLOR_MUTED)

    name      = review.get('name', '').strip()
    field_rev = review.get('field', '').strip()
    stars     = review.get('stars', 0)
    comment   = review.get('comment', '').strip()

    # 별점 표시
    star_str = '★' * stars + '☆' * (5 - stars)

    # 한 줄 요약 헤더 (이름, 분야, 별점)
    p = doc.add_paragraph()
    if name:
        r = p.add_run(name)
        sf(r, bold=True, size=11, color=COLOR_SECONDARY)
    if field_rev:
        r = p.add_run(f'  |  {field_rev}')
        sf(r, size=10, color=COLOR_MUTED)
    if stars:
        r = p.add_run(f'  {star_str}  ({stars}/5)')
        c = COLOR_STARS_HIGH if stars >= 4 else COLOR_STARS_LOW
        sf(r, bold=True, size=11, color=c)
    p.paragraph_format.space_after = Pt(6)

    # 한줄평
    if comment:
        p = doc.add_paragraph()
        r = p.add_run(f'"{comment}"')
        sf(r, size=10, color=COLOR_SECONDARY)
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(4)
